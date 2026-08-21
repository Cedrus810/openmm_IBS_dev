from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "16324F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
RULE = "B8C4D1"
WHITE = "FFFFFF"


def set_run_font(run, latin="Times New Roman", east_asia="Times New Roman", size=12, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        insert_before = None
        for tag in ("w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark"):
            insert_before = tc_pr.find(qn(tag))
            if insert_before is not None:
                break
        if insert_before is None:
            tc_pr.append(shd)
        else:
            insert_before.addprevious(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    # OfficeCLI 1.0.144 rejects python-docx's cell-level tcMar ordering.
    # Keep the stable document defaults instead of emitting non-portable XML.
    return


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_fixed_table_geometry(table, widths_twips):
    tbl_pr = table._tbl.tblPr
    table.autofit = False
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.insert_element_before(
            tbl_ind,
            "w:tblBorders",
            "w:shd",
            "w:tblLayout",
            "w:tblCellMar",
            "w:tblLook",
        )
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_twips[i]))
            tc_w.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction, cached="1"):
    begin = paragraph.add_run()
    begin._r.append(OxmlElement("w:fldChar"))
    begin._r[-1].set(qn("w:fldCharType"), "begin")
    instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    instr._r.append(instr_text)
    separate = paragraph.add_run()
    separate._r.append(OxmlElement("w:fldChar"))
    separate._r[-1].set(qn("w:fldCharType"), "separate")
    value = paragraph.add_run(cached)
    end = paragraph.add_run()
    end._r.append(OxmlElement("w:fldChar"))
    end._r[-1].set(qn("w:fldCharType"), "end")
    return value


def add_toc(paragraph):
    add_field(paragraph, 'TOC \\o "1-4" \\h \\z \\u', "Right-click and update field")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(12)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Times New Roman"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string(MID_GRAY)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)

    heading_specs = {
        "Heading 1": (12, BLUE, 18, 10),
        "Heading 2": (12, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
        "Heading 4": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        st = styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True

    for name, base in (("Code Block", "Normal"), ("Language Label", "Normal"), ("Equation Placeholder", "Normal"), ("Cover Meta", "Normal")):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st = styles[name]
        st.base_style = styles[base]

    code = styles["Code Block"]
    code.font.name = "Times New Roman"
    code.font.size = Pt(12)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.0

    language = styles["Language Label"]
    language.font.name = "Times New Roman"
    language.font.size = Pt(12)
    language.font.bold = True
    language.font.color.rgb = RGBColor.from_string(BLUE)
    language._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    language.paragraph_format.space_before = Pt(5)
    language.paragraph_format.space_after = Pt(3)
    language.paragraph_format.keep_with_next = True

    eq = styles["Equation Placeholder"]
    eq.font.name = "Times New Roman"
    eq.font.size = Pt(12)
    eq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.space_before = Pt(5)
    eq.paragraph_format.space_after = Pt(7)
    eq.paragraph_format.keep_together = True

    meta = styles["Cover Meta"]
    meta.font.name = "Times New Roman"
    meta.font.size = Pt(12)
    meta.font.color.rgb = RGBColor.from_string(MID_GRAY)
    meta._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)


INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\\\([^)]*\\\))")


def clean_inline_text(text):
    text = text.replace("  ", " ")
    text = text.replace("\\|", "|")
    return text


def inline_math_to_text(formula):
    """Turn short inline LaTeX into readable linear math, never raw commands."""
    text = formula
    command_map = {
        r"\Delta": "Δ",
        r"\delta": "δ",
        r"\theta": "θ",
        r"\phi": "φ",
        r"\lambda": "λ",
        r"\beta": "β",
        r"\alpha": "α",
        r"\sigma": "σ",
        r"\pi": "π",
        r"\langle": "⟨",
        r"\rangle": "⟩",
        r"\neq": "≠",
        r"\simeq": "≈",
        r"\propto": "∝",
        r"\times": "×",
        r"\circ": "°",
        r"\rightarrow": "→",
        r"\to": "→",
    }
    for source, target in command_map.items():
        text = text.replace(source, target)
    # Common one-argument styling commands are unnecessary in linear prose.
    previous = None
    while previous != text:
        previous = text
        text = re.sub(r"\\(?:mathbf|mathrm|mathit|mathcal|operatorname|text)\{([^{}]*)\}", r"\1", text)
    text = text.replace(r"\left", "").replace(r"\right", "")
    text = text.replace(r"\,", " ").replace(r"\;", " ")
    text = re.sub(r"_\{([^{}]+)\}", r"_\1", text)
    text = re.sub(r"\^\{([^{}]+)\}", r"^\1", text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace("\\", "")
    return text.strip()


def add_inline(paragraph, text, *, base_size=12, color=None):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(clean_inline_text(text[pos:match.start()]))
            set_run_font(run, size=base_size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            set_run_font(run, size=base_size, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=12, color=DARK_BLUE)
        else:
            formula = token[2:-2]
            run = paragraph.add_run(inline_math_to_text(formula))
            set_run_font(run, size=12, color=color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(clean_inline_text(text[pos:]))
        set_run_font(run, size=base_size, color=color)


def split_table_row(line):
    stripped = line.strip().strip("|")
    return [c.strip() for c in re.split(r"(?<!\\)\|", stripped)]


def is_separator_row(cells):
    return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells)


def add_markdown_table(doc, rows):
    parsed = [split_table_row(r) for r in rows]
    if len(parsed) > 1 and is_separator_row(parsed[1]):
        parsed.pop(1)
    cols = max(len(r) for r in parsed)
    for r in parsed:
        r.extend([""] * (cols - len(r)))
    table = doc.add_table(rows=len(parsed), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    if cols == 2:
        widths = [2700, 6660]
    elif cols == 3:
        widths = [2100, 3630, 3630]
    elif cols == 4:
        widths = [1750, 2535, 2535, 2540]
    elif cols == 5:
        widths = [1500, 1965, 1965, 1965, 1965]
    else:
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)
    set_fixed_table_geometry(table, widths)
    for ri, values in enumerate(parsed):
        row = table.rows[ri]
        prevent_row_split(row)
        if ri == 0:
            set_repeat_table_header(row)
        for ci, value in enumerate(values):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if ri == 0:
                set_cell_shading(cell, PALE_BLUE)
            elif ri % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, base_size=12)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    v = add_field(p, "PAGE", "1")
    set_run_font(v, size=12, color=MID_GRAY)


def add_running_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("ABFE-IBS • METHODS, VALIDATION, AND DECISION RECORD")
    set_run_font(r, size=12, color=MID_GRAY)


def normalize_formula(formula):
    formula = formula.strip()
    formula = re.sub(r"\\begin\{(?:aligned|gathered|split)\}", "", formula)
    formula = re.sub(r"\\end\{(?:aligned|gathered|split)\}", "", formula)
    formula = formula.replace("&=", "=").replace("&", "")
    formula = formula.replace("\\\\", " \\quad ")
    return " ".join(line.strip() for line in formula.splitlines() if line.strip())


def build(md_path: Path, out_path: Path, manifest_path: Path, lang: str):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    section.different_first_page_header_footer = True
    add_page_number_footer(section)

    # Cover
    p = doc.add_paragraph(style="Cover Meta")
    p.paragraph_format.space_before = Pt(32)
    cover_kicker = "导师技术报告" if lang == "zh" else "ADVISOR-FACING TECHNICAL REPORT"
    r = p.add_run(cover_kicker)
    r.bold = True
    set_run_font(r, size=12, color=BLUE)
    cover_title = "ABFE-IBS 技术报告" if lang == "zh" else "ABFE-IBS Technical Report"
    cover_subtitle = "从物理路径、采样分布到当前证据边界" if lang == "zh" else "From the Physical Path and Sampling Distribution to the Current Evidence Boundary"
    date_line = "整理日期：2026-08-20" if lang == "zh" else "Consolidation date: 2026-08-20"
    source_line = "内容基线：代码、机器可读结果、日志、checkpoint 与预注册协议" if lang == "zh" else "Evidence basis: code, machine-readable results, logs, checkpoints, and preregistration protocols"
    p = doc.add_paragraph(style="Title")
    add_inline(p, cover_title, base_size=12, color=NAVY)
    p = doc.add_paragraph(style="Subtitle")
    add_inline(p, cover_subtitle, base_size=12, color=MID_GRAY)
    p = doc.add_paragraph(style="Cover Meta")
    add_inline(p, date_line, base_size=12, color=MID_GRAY)
    p = doc.add_paragraph(style="Cover Meta")
    add_inline(p, source_line, base_size=12, color=MID_GRAY)
    p.paragraph_format.space_after = Pt(24)

    box = doc.add_table(rows=3, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.style = "Table Grid"
    set_fixed_table_geometry(box, [9360])
    cover_notes = (
        [
            "目的：说明已完成的工作、当前证据、尚未闭合的风险与下一阶段实验。",
            "主线：物理输入 → Boresch → PME → ACE/IBS → actual-mixture estimator → 当前结论。",
            "证据规则：可复核 artifact 与代码事实优先于叙述性判断。",
        ]
        if lang == "zh"
        else [
            "Purpose: explain completed work, current evidence, unresolved risks, and the next experiments.",
            "Main line: physical input → Boresch → PME → ACE/IBS → actual-mixture estimator → current conclusion.",
            "Evidence rule: auditable artifacts and code-level facts take precedence over narrative judgment.",
        ]
    )
    for idx, note in enumerate(cover_notes):
        cell = box.cell(idx, 0)
        set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
        set_cell_shading(cell, "F7FAFC" if idx % 2 == 0 else WHITE)
        p = cell.paragraphs[0]
        add_inline(p, note, base_size=12)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    p = doc.add_paragraph("目录" if lang == "zh" else "Contents", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    toc = doc.add_paragraph()
    add_toc(toc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    equations = []
    i = 0
    # Skip the original two H1 titles and cover blockquotes until the first rule.
    while i < len(lines) and lines[i].strip() != "---":
        i += 1
    if i < len(lines):
        i += 1

    while i < len(lines):
        raw = lines[i]
        s = raw.strip()
        if not s:
            i += 1
            continue
        if s == "---":
            i += 1
            continue
        if s.startswith("\\["):
            formula_lines = []
            after_open = s[2:].strip()
            if after_open:
                formula_lines.append(after_open)
            i += 1
            while i < len(lines) and "\\]" not in lines[i]:
                formula_lines.append(lines[i])
                i += 1
            if i < len(lines):
                before_close = lines[i].split("\\]", 1)[0].strip()
                if before_close:
                    formula_lines.append(before_close)
            formula = normalize_formula("\n".join(formula_lines))
            token = f"[[EQ:{len(equations)+1:04d}]]"
            p = doc.add_paragraph(style="Equation Placeholder")
            p.add_run(token)
            equations.append({"token": token, "formula": formula})
            i += 1
            continue
        if s.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph(style="Code Block")
            run = p.add_run("\n".join(code_lines))
            set_run_font(run, size=12)
            i += 1
            continue
        if s.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = [raw]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", s)
        if heading:
            level = min(max(len(heading.group(1)) - 1, 1), 4)
            text = heading.group(2).strip()
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, text, base_size=12, color=BLUE if level < 3 else DARK_BLUE)
            i += 1
            continue
        if s in {"**中文**", "**English**", "**中文说明**", "**English notes**", "**中文结论**", "**English conclusion**"}:
            p = doc.add_paragraph(style="Language Label")
            add_inline(p, s.strip("*"), base_size=12, color=BLUE)
            i += 1
            continue
        if s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.right_indent = Inches(0.12)
            p.paragraph_format.space_after = Pt(5)
            add_inline(p, s.lstrip("> "), base_size=12, color=MID_GRAY)
            i += 1
            continue
        bullet = re.match(r"^[-*+]\s+(.*)$", s)
        numbered = re.match(r"^\d+[.)]\s+(.*)$", s)
        if bullet or numbered:
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            add_inline(p, (bullet or numbered).group(1), base_size=12)
            i += 1
            continue

        # Join wrapped Markdown lines into a single human-readable paragraph.
        para_lines = [s]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt == "---" or nxt.startswith(("#", "|", ">", "```", "\\[")):
                break
            if re.match(r"^[-*+]\s+", nxt) or re.match(r"^\d+[.)]\s+", nxt):
                break
            para_lines.append(nxt)
            i += 1
        text = " ".join(para_lines)
        p = doc.add_paragraph()
        add_inline(p, text, base_size=12)

    props = doc.core_properties
    props.title = "ABFE-IBS Technical Documentation"
    props.subject = "Methods, results, validation, failed routes, and next-stage decisions"
    props.author = "ABFE-IBS Project"
    props.keywords = "ABFE, IBS, Boresch, PME, ACE, MBAR, MACE, free energy"
    props.comments = "Rebuilt from the bilingual Markdown technical record; equations are finalized with OfficeCLI."

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    manifest_path.write_text(json.dumps(equations, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"output": str(out_path), "equations": len(equations), "manifest": str(manifest_path)}, ensure_ascii=False))


if __name__ == "__main__":
    if len(sys.argv) not in (4, 5):
        raise SystemExit("usage: build_revised_docx.py INPUT.md OUTPUT.docx EQUATIONS.json [zh|en]")
    lang = sys.argv[4] if len(sys.argv) == 5 else "zh"
    if lang not in {"zh", "en"}:
        raise SystemExit("language must be zh or en")
    build(Path(sys.argv[1]), Path(sys.argv[2]), Path(sys.argv[3]), lang)
